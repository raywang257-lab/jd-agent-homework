"""多格式导出 -- Word / Markdown / 纯文本 / PDF"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_jd_to_docx(jd_text: str) -> bytes:
    """将 Markdown JD 转换为 Word 文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    lines = jd_text.split("\n")
    for line in lines:
        line = line.rstrip()

        if not line:
            doc.add_paragraph("")
            continue

        # 标题
        if line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            # 处理行内粗体
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part.strip("*"))
                    run.bold = True
                elif part:
                    p.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def jd_to_markdown(jd_text: str) -> str:
    """直接返回 Markdown 格式"""
    return jd_text


def jd_to_plain_text(jd_text: str) -> str:
    """将 Markdown JD 转换为纯文本"""
    text = jd_text
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"^- ", "• ", text, flags=re.MULTILINE)
    return text


def jd_to_pdf(jd_text: str) -> bytes:
    """将 JD 转换为 PDF（使用 reportlab）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            ListFlowable,
            ListItem,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise ValueError("reportlab 未安装，无法导出 PDF。请运行 pip install reportlab")

    buffer = io.BytesIO()

    # 尝试注册中文字体
    font_name = "Helvetica"
    font_paths = [
        ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
        ("SimHei", "C:/Windows/Fonts/simhei.ttf"),
        ("MSYH", "C:/Windows/Fonts/msyh.ttc"),
    ]
    for name, path in font_paths:
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            font_name = name
            break
        except Exception:
            continue

    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        spaceAfter=6,
        spaceBefore=12,
    )
    normal_style = ParagraphStyle(
        "CustomNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=18,
    )
    bullet_style = ParagraphStyle(
        "CustomBullet",
        parent=normal_style,
        leftIndent=20,
    )

    story: list = []

    lines = jd_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue

        if line.startswith("# "):
            story.append(Paragraph(line[2:].strip(), heading_style))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:].strip(), h2_style))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), h2_style))
        elif line.startswith("- "):
            story.append(Paragraph("• " + line[2:].strip(), bullet_style))
        elif line.startswith("**") and line.endswith("**"):
            bold_text = line.strip("*")
            story.append(Paragraph(f"<b>{bold_text}</b>", normal_style))
        else:
            # 处理行内粗体
            html_line = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", line)
            story.append(Paragraph(html_line, normal_style))

    doc.build(story)
    return buffer.getvalue()
